import pymupdf
import pytest
from docx import Document

from app.services.parsers.docx_parser import DocxParser
from app.services.parsers.parser_factory import get_parser
from app.services.parsers.pdf_parser import PdfParser
from app.services.parsers.txt_parser import TxtParser


def test_docx_parser_extracts_paragraphs(tmp_path):
    file_path = tmp_path / "policy.docx"
    document = Document()
    document.add_paragraph("Annual leave policy")
    document.add_paragraph("")
    document.add_paragraph("Employees receive paid leave.")
    document.save(file_path)

    content = DocxParser().parse(file_path)

    assert content == "Annual leave policy\nEmployees receive paid leave."


def test_txt_parser_extracts_utf8_text(tmp_path):
    file_path = tmp_path / "guide.txt"
    file_path.write_text("新员工入职指南\n请准备身份证明。", encoding="utf-8")

    assert TxtParser().parse(file_path) == "新员工入职指南\n请准备身份证明。"


def test_pdf_parser_extracts_page_text(tmp_path):
    file_path = tmp_path / "manual.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Enterprise knowledge manual")
    document.save(file_path)
    document.close()

    content = PdfParser().parse(file_path)

    assert "Enterprise knowledge manual" in content


def test_parser_factory_rejects_unsupported_format(tmp_path):
    with pytest.raises(ValueError, match="不支持的文件格式"):
        get_parser(tmp_path / "archive.zip")


@pytest.mark.parametrize(
    ("filename", "parser_type"),
    [("a.docx", DocxParser), ("a.txt", TxtParser), ("a.pdf", PdfParser)],
)
def test_parser_factory_selects_by_extension(filename, parser_type):
    assert isinstance(get_parser(filename), parser_type)
